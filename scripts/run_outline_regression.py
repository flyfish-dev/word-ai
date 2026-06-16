from __future__ import annotations

import shutil
import zipfile
from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document
from lxml import etree

from word_ai_mcp.ooxml import get_outline, inspect_docx, list_anchors, list_paragraphs, read_heading_section

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W_NS}


def qn(tag: str) -> str:
    prefix, name = tag.split(":", 1)
    ns = {"w": W_NS}[prefix]
    return f"{{{ns}}}{name}"


def ensure_ppr(p: etree._Element) -> etree._Element:
    ppr = p.find("./w:pPr", namespaces=NS)
    if ppr is None:
        ppr = etree.Element(qn("w:pPr"))
        p.insert(0, ppr)
    return ppr


def set_style(p: etree._Element, style_id: str) -> None:
    ppr = ensure_ppr(p)
    pstyle = ppr.find("./w:pStyle", namespaces=NS)
    if pstyle is None:
        pstyle = etree.Element(qn("w:pStyle"))
        ppr.insert(0, pstyle)
    pstyle.set(qn("w:val"), style_id)


def set_direct_outline(p: etree._Element, value: str) -> None:
    ppr = ensure_ppr(p)
    outline = ppr.find("./w:outlineLvl", namespaces=NS)
    if outline is None:
        outline = etree.SubElement(ppr, qn("w:outlineLvl"))
    outline.set(qn("w:val"), value)


def clear_runs(p: etree._Element) -> None:
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def add_run_text(p: etree._Element, text: str) -> None:
    r = etree.SubElement(p, qn("w:r"))
    t = etree.SubElement(r, qn("w:t"))
    t.text = text


def add_fld_char(p: etree._Element, kind: str) -> None:
    r = etree.SubElement(p, qn("w:r"))
    fld = etree.SubElement(r, qn("w:fldChar"))
    fld.set(qn("w:fldCharType"), kind)


def add_instr(p: etree._Element, instr: str) -> None:
    r = etree.SubElement(p, qn("w:r"))
    node = etree.SubElement(r, qn("w:instrText"))
    node.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    node.text = instr


def add_paragraph_style(styles_root: etree._Element, style_id: str, name: str, outline_level: str | None = None) -> None:
    old = styles_root.xpath("//w:style[@w:styleId=$sid]", sid=style_id, namespaces=NS)
    for item in old:
        item.getparent().remove(item)
    style = etree.SubElement(styles_root, qn("w:style"))
    style.set(qn("w:type"), "paragraph")
    style.set(qn("w:styleId"), style_id)
    name_el = etree.SubElement(style, qn("w:name"))
    name_el.set(qn("w:val"), name)
    if outline_level is not None:
        ppr = etree.SubElement(style, qn("w:pPr"))
        outline = etree.SubElement(ppr, qn("w:outlineLvl"))
        outline.set(qn("w:val"), outline_level)


def rewrite_fixture(path: Path, close_toc: bool = True) -> None:
    with zipfile.ZipFile(path, "r") as zin:
        entries = {i.filename: zin.read(i.filename) for i in zin.infolist() if not i.is_dir()}
        infos = {i.filename: i for i in zin.infolist() if not i.is_dir()}

    parser = etree.XMLParser(remove_blank_text=False)
    styles_root = etree.fromstring(entries["word/styles.xml"], parser)
    for style_id, name, outline in [
        ("1", "heading 1", "0"),
        ("2", "heading 2", "1"),
        ("CN1", "标题1", None),
        ("CN2", "标题 2", None),
        ("APP1", "附录一级标题", "0"),
        ("TOCHeading", "TOC Heading", "9"),
        ("TOC1", "toc 1", None),
        ("TOC2", "目录 2", None),
    ]:
        add_paragraph_style(styles_root, style_id, name, outline)

    doc_root = etree.fromstring(entries["word/document.xml"], parser)
    paras = doc_root.xpath("//w:body//w:p", namespaces=NS)
    styles = [
        None,
        "TOCHeading",
        "TOC1",
        "TOC2",
        "TOC1",
        "1",
        "2",
        "CN1",
        "CN2",
        None,
        "APP1",
    ]
    for p, style in zip(paras, styles):
        if style:
            set_style(p, style)

    clear_runs(paras[2])
    add_fld_char(paras[2], "begin")
    add_instr(paras[2], ' TOC \\o "1-3" \\h \\z \\u ')
    add_fld_char(paras[2], "separate")
    add_run_text(paras[2], "1 功能介绍3")

    clear_runs(paras[4])
    add_run_text(paras[4], "2 使用指南4")
    if close_toc:
        add_fld_char(paras[4], "end")

    set_direct_outline(paras[9], "2")

    entries["word/styles.xml"] = etree.tostring(styles_root, encoding="UTF-8", xml_declaration=True, standalone=True)
    entries["word/document.xml"] = etree.tostring(doc_root, encoding="UTF-8", xml_declaration=True, standalone=True)

    tmp = path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(tmp, "w") as zout:
        for name, data in entries.items():
            zout.writestr(infos[name], data)
    shutil.move(tmp, path)


def make_fixture(path: Path, close_toc: bool = True) -> None:
    doc = Document()
    for text in [
        "中文大纲识别回归",
        "目录",
        "1 功能介绍3",
        "1.1 系统概述3",
        "2 使用指南4",
        "功能介绍",
        "目录管理系统",
        "中文一级标题",
        "中文二级标题",
        "直接大纲级别",
        "附录一",
    ]:
        doc.add_paragraph(text)
    doc.save(path)
    rewrite_fixture(path, close_toc=close_toc)


def assert_outline_fixture(docx_path: Path) -> None:
    outline = get_outline(docx_path)
    headings = outline["headings"]
    texts = [h["text"] for h in headings]
    assert texts == ["功能介绍", "目录管理系统", "中文一级标题", "中文二级标题", "直接大纲级别", "附录一"], texts
    assert [h["level"] for h in headings] == [1, 2, 1, 2, 3, 1], headings
    assert all(not h["is_toc"] for h in headings), headings

    paragraphs = list_paragraphs(docx_path, include_empty=False)["paragraphs"]
    toc_items = [p for p in paragraphs if p["is_toc"]]
    assert [p["text_preview"] for p in toc_items] == ["目录", "1 功能介绍3", "1.1 系统概述3", "2 使用指南4"], toc_items
    assert all(p["heading_level"] is None for p in toc_items), toc_items

    anchors = [a.to_dict() for a in list_anchors(docx_path)]
    heading_anchor_texts = [a["text_preview"] for a in anchors if a["kind"] == "heading"]
    assert heading_anchor_texts == texts, heading_anchor_texts
    assert not any((a.get("text_preview") or "") in {"目录", "1 功能介绍3", "1.1 系统概述3", "2 使用指南4"} for a in anchors if a["kind"] == "heading"), anchors

    section = read_heading_section(docx_path, heading_text="功能介绍")
    assert section["heading"]["style_name"] == "heading 1"
    assert "1 功能介绍3" not in section["text"]

    info = inspect_docx(docx_path)
    assert info["heading_count"] == 6, info


def main() -> int:
    with TemporaryDirectory() as td:
        closed_docx = Path(td) / "outline-regression.docx"
        make_fixture(closed_docx)
        assert_outline_fixture(closed_docx)

        unclosed_docx = Path(td) / "outline-regression-unclosed-toc.docx"
        make_fixture(unclosed_docx, close_toc=False)
        assert_outline_fixture(unclosed_docx)

    print("outline regression passed")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
