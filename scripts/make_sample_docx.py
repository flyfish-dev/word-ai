from __future__ import annotations

import shutil
import zipfile
from pathlib import Path
from random import randint

from docx import Document
from docx.enum.section import WD_SECTION
from docx.shared import Inches
from lxml import etree

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W_NS}

def qn(tag: str) -> str:
    prefix, name = tag.split(":", 1)
    ns = {"w": W_NS}[prefix]
    return f"{{{ns}}}{name}"


def wrap_paragraphs_with_content_controls(docx_path: Path, marker_to_tag: dict[str, tuple[str, str]]) -> None:
    with zipfile.ZipFile(docx_path, "r") as zin:
        entries = {i.filename: zin.read(i.filename) for i in zin.infolist() if not i.is_dir()}
        infos = {i.filename: i for i in zin.infolist() if not i.is_dir()}
    parser = etree.XMLParser(remove_blank_text=False)
    root = etree.fromstring(entries["word/document.xml"], parser)
    for marker, (tag, alias) in marker_to_tag.items():
        for p in root.xpath("//w:body//w:p", namespaces=NS):
            text = "".join(t.text or "" for t in p.xpath(".//w:t", namespaces=NS))
            if marker in text:
                parent = p.getparent()
                idx = parent.index(p)
                parent.remove(p)
                sdt = etree.Element(qn("w:sdt"))
                sdt_pr = etree.SubElement(sdt, qn("w:sdtPr"))
                alias_el = etree.SubElement(sdt_pr, qn("w:alias"))
                alias_el.set(qn("w:val"), alias)
                tag_el = etree.SubElement(sdt_pr, qn("w:tag"))
                tag_el.set(qn("w:val"), tag)
                id_el = etree.SubElement(sdt_pr, qn("w:id"))
                id_el.set(qn("w:val"), str(randint(100000, 999999)))
                lock = etree.SubElement(sdt_pr, qn("w:lock"))
                lock.set(qn("w:val"), "sdtLocked")
                content = etree.SubElement(sdt, qn("w:sdtContent"))
                content.append(p)
                parent.insert(idx, sdt)
                break
    entries["word/document.xml"] = etree.tostring(root, encoding="UTF-8", xml_declaration=True, standalone=True)
    tmp = docx_path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(tmp, "w") as zout:
        for name, data in entries.items():
            info = infos[name]
            zout.writestr(info, data)
    shutil.move(tmp, docx_path)


def main():
    out = Path(__file__).resolve().parent.parent / "examples" / "sample_contract.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.header.paragraphs[0].text = "项目交付文档模板 | Word AI MCP Sample"
    section.footer.paragraphs[0].text = "机密 | 示例文件 | Page field should be refreshed in Word"

    doc.add_heading("需求规格说明书", 0)
    doc.add_paragraph("文档编号：SRS-DEMO-001    版本：V1.0")

    doc.add_heading("1. 引言", level=1)
    doc.add_paragraph("[[CC:overview]] 本文档描述系统建设目标、范围、角色和总体约束。该段落由内容控件锚定，可由 AI 安全替换。")
    doc.add_heading("1.1 项目背景", level=2)
    doc.add_paragraph("本项目面向多部门协同场景，要求文档生成流程可审计、可回滚、可复用。")

    doc.add_heading("2. 功能需求", level=1)
    doc.add_paragraph("[[CC:functional_scope]] 系统需支持需求录入、需求评审、版本追踪和交付文档自动生成。")
    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    table.cell(0, 0).text = "编号"
    table.cell(0, 1).text = "功能"
    table.cell(0, 2).text = "说明"
    table.cell(1, 0).text = "FR-001"
    table.cell(1, 1).text = "文档索引"
    table.cell(1, 2).text = "解析标题、表格、内容控件和字段。"
    table.cell(2, 0).text = "FR-002"
    table.cell(2, 1).text = "增量编辑"
    table.cell(2, 2).text = "仅修改指定锚点范围。"
    table.cell(3, 0).text = "FR-003"
    table.cell(3, 1).text = "审计回滚"
    table.cell(3, 2).text = "保留修改前后版本和差异报告。"

    doc.add_heading("3. 非功能需求", level=1)
    doc.add_paragraph("[[CC:nfr_performance]] 超大文档处理应使用索引、分块读取和定点回写，避免全量重建。")
    doc.add_paragraph("安全要求：所有写操作必须默认生成备份和审计记录。")

    doc.add_heading("4. 数据库设计", level=1)
    doc.add_paragraph("数据库设计章节用于展示表格和字段说明的稳定保留。")
    db = doc.add_table(rows=3, cols=4)
    db.style = "Table Grid"
    for i, h in enumerate(["字段", "类型", "约束", "说明"]):
        db.cell(0, i).text = h
    db.cell(1, 0).text = "doc_id"
    db.cell(1, 1).text = "varchar(64)"
    db.cell(1, 2).text = "PK"
    db.cell(1, 3).text = "文档唯一标识"
    db.cell(2, 0).text = "anchor_tag"
    db.cell(2, 1).text = "varchar(128)"
    db.cell(2, 2).text = "INDEX"
    db.cell(2, 3).text = "内容控件锚点"

    doc.save(out)
    wrap_paragraphs_with_content_controls(
        out,
        {
            "[[CC:overview]]": ("WORD-AI:SRS:1.0:overview", "概述正文锚点"),
            "[[CC:functional_scope]]": ("WORD-AI:SRS:2.0:functional_scope", "功能需求范围锚点"),
            "[[CC:nfr_performance]]": ("WORD-AI:SRS:3.0:nfr_performance", "性能需求锚点"),
        },
    )
    print(out)


if __name__ == "__main__":
    main()
